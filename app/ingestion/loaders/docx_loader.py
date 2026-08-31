"""
Word documents via python-docx. Extracts paragraph text and table cells
(tables are flattened to " | "-joined rows and appended after the body
text, tagged with a header, rather than interleaved in original position
— good enough for retrieval, not a layout-preserving conversion).
"""
from __future__ import annotations

from pathlib import Path

DOCX_EXTENSIONS = {".docx"}


def load_docx(path: Path) -> tuple[str, dict]:
    from docx import Document

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    table_lines: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            table_lines.append(" | ".join(cells))

    parts = list(paragraphs)
    if table_lines:
        parts += ["", "--- Tables ---", *table_lines]
    text = "\n".join(parts)

    metadata = {
        "source_path": str(path),
        "filename": path.name,
        "format": "docx",
        "num_paragraphs": len(paragraphs),
        "num_tables": len(doc.tables),
    }
    return text, metadata
